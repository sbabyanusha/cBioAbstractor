"""
cbioportal_curator.py
─────────────────────
Core curation engine for the Synopsis tool.

Given an uploaded paper PDF and one or more supplementary Excel files, this
module:
  1. Extracts study metadata from the PDF via the RAG vector store.
  2. Inspects every supplementary Excel sheet and classifies it against the
     cBioPortal file-format catalogue.
  3. Produces per-file curation instructions (column mappings, required
     transformations, data gaps).
  4. Emits a fully formatted .docx curation report.

Public API
──────────
    curate(pdf_path, supp_paths, llm_model="openai/gpt-4o",
           temperature=0.2) -> dict
        Returns {"report_path": str, "summary": dict}
"""

from __future__ import annotations

import io
import os
import re
import tempfile
from pathlib import Path
from typing import Any

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import load_chat_model
from langchain_core.messages import HumanMessage, SystemMessage
from spec_match import classify_sheet, ClassificationResult
from cbioportal_spec import SPEC_BY_KEY

# ─────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────

CURABILITY = {
    "CLINICAL_PATIENT":   ("YES",     "HIGH"),
    "CLINICAL_SAMPLE":    ("YES",     "HIGH"),
    "MUTATION_MAF":       ("PARTIAL", "HIGH"),
    "STRUCTURAL_VARIANT": ("YES",     "HIGH"),
    "DISCRETE_CNA":       ("PARTIAL", "MEDIUM"),
    "CONTINUOUS_CNA":     ("PARTIAL", "MEDIUM"),
    "SEGMENTED":          ("PARTIAL", "MEDIUM"),
    "EXPRESSION":         ("PARTIAL", "MEDIUM"),
    "METHYLATION":        ("PARTIAL", "LOW"),
    "MUTSIG":             ("PARTIAL", "MEDIUM"),
    "GISTIC":             ("PARTIAL", "MEDIUM"),
    "GENERIC_ASSAY":      ("PARTIAL", "LOW"),
    "NOT_LOADABLE":       ("NO",      "N/A"),
}

# MAF Variant_Classification remapping from ANNOVAR terms
ANNOVAR_TO_MAF = {
    "nonsynonymous snv":         "Missense_Mutation",
    "synonymous snv":            "Silent",
    "stopgain":                  "Nonsense_Mutation",
    "stoploss":                  "Nonstop_Mutation",
    "frameshift deletion":       "Frame_Shift_Del",
    "frameshift insertion":      "Frame_Shift_Ins",
    "nonframeshift deletion":    "In_Frame_Del",
    "nonframeshift insertion":   "In_Frame_Ins",
    "splicing":                  "Splice_Site",
    "unknown":                   "Unknown",
}

SYSTEM_PROMPT_CURATOR = """
You are an expert bioinformatics data curator specialising in the cBioPortal
platform (https://docs.cbioportal.org/file-formats/).

When given text extracted from a cancer genomics paper, extract the following
study metadata and return it as a JSON object with exactly these keys:

{
  "study_title": "...",
  "cancer_type": "...",           // short abbreviation e.g. brca, gist, luad
  "cancer_type_full": "...",      // e.g. Breast Invasive Carcinoma
  "num_samples": "...",           // integer or string
  "num_patients": "...",          // integer or string
  "reference_genome": "...",      // hg19 or hg38
  "sequencing_types": ["..."],    // e.g. ["WES","WGS","WTS"]
  "pmid": "...",                  // PubMed ID if mentioned
  "doi": "...",                   // DOI string
  "first_author_surname": "...",
  "year": "...",
  "journal": "...",
  "study_id_suggestion": "...",   // snake_case e.g. gist_xie_2024
  "description": "...",          // one sentence
  "key_findings": ["..."],        // up to 5 bullet points
  "primary_site": "...",          // anatomical site e.g. "Stomach and small intestine"
  "cohort_description": "...",    // one sentence describing the cohort composition
  "meta_description": "...",      // concise description for meta_study.txt (200 chars max)
  "data_repositories": ["..."],   // GEO/GDC/SRA accession strings mentioned in paper
  "corresponding_authors": "..."  // name and email of corresponding authors if mentioned
}

Return ONLY the JSON — no markdown fences, no extra text.
"""

# ─────────────────────────────────────────────────────────────
# PDF helpers
# ─────────────────────────────────────────────────────────────

def _extract_pdf_text(pdf_path: str, max_pages: int = 12) -> str:
    reader = PdfReader(pdf_path)
    pages = reader.pages[:max_pages]
    return "\n".join(p.extract_text() or "" for p in pages)


# ─────────────────────────────────────────────────────────────
# Excel helpers
# ─────────────────────────────────────────────────────────────

def _read_excel_sheets(path: str) -> dict[str, pd.DataFrame]:
    """Return all sheets as DataFrames, stripping blank leading rows."""
    xl = pd.ExcelFile(path)
    sheets: dict[str, pd.DataFrame] = {}
    for name in xl.sheet_names:
        df = xl.parse(name, header=None)
        df = df.dropna(how="all")
        sheets[name] = df
    return sheets


def _read_file_as_sheets(path: str) -> dict[str, pd.DataFrame]:
    """
    Universal file reader — returns a dict of {sheet_name: DataFrame}
    for any supported supplementary file format:

      .xlsx / .xls          — Excel (multi-sheet)
      .csv                  — comma-separated
      .tsv / .tab           — tab-separated
      .txt                  — auto-detected delimiter (tab > comma > pipe > space)
      .doc / .docx          — Word document: tables extracted as DataFrames,
                              plain text paragraphs as a single-column DataFrame
      .maf                  — tab-separated MAF (treated as .tsv)
    """
    from pathlib import Path as _Path
    import io as _io

    ext = _Path(path).suffix.lower()

    # ── Excel ────────────────────────────────────────────────────────────
    if ext in (".xlsx", ".xls"):
        return _read_excel_sheets(path)

    # ── CSV ──────────────────────────────────────────────────────────────
    if ext == ".csv":
        df = pd.read_csv(path, header=None, dtype=str, encoding_errors="replace")
        df = df.dropna(how="all")
        return {"Sheet1": df}

    # ── TSV / TAB / MAF ──────────────────────────────────────────────────
    if ext in (".tsv", ".tab", ".maf"):
        df = pd.read_csv(path, sep="	", header=None, dtype=str, encoding_errors="replace")
        df = df.dropna(how="all")
        return {"Sheet1": df}

    # ── TXT — sniff delimiter ─────────────────────────────────────────────
    if ext == ".txt":
        raw = open(path, "r", encoding="utf-8", errors="replace").read(4096)
        # count candidate delimiters
        counts = {"	": raw.count("	"), ",": raw.count(","),
                  "|": raw.count("|"),   " ": raw.count(" ")}
        sep = max(counts, key=counts.get)
        if counts[sep] == 0:
            sep = "	"   # fallback
        df = pd.read_csv(path, sep=sep, header=None, dtype=str,
                         encoding_errors="replace", on_bad_lines="skip")
        df = df.dropna(how="all")
        return {"Sheet1": df}

    # ── DOC / DOCX ────────────────────────────────────────────────────────
    if ext in (".doc", ".docx"):
        try:
            from docx import Document as _DocxDoc
        except ImportError:
            raise ImportError("python-docx is required to read .doc/.docx files. "
                              "Install with: pip install python-docx")

        if ext == ".doc":
            # Legacy .doc: try to convert via LibreOffice, else read as text
            import subprocess, tempfile, shutil
            if shutil.which("libreoffice"):
                with tempfile.TemporaryDirectory() as tmp:
                    subprocess.run(
                        ["libreoffice", "--headless", "--convert-to", "docx",
                         "--outdir", tmp, path],
                        capture_output=True, timeout=30
                    )
                    converted = list(_Path(tmp).glob("*.docx"))
                    if converted:
                        path = str(converted[0])
                        ext  = ".docx"
                    else:
                        # fall through to plain-text extraction
                        ext = "_unknown"
            else:
                ext = "_unknown"   # no converter available

        if ext == ".docx":
            doc   = _DocxDoc(path)
            result: dict[str, pd.DataFrame] = {}

            # Extract each table as its own sheet
            for t_idx, tbl in enumerate(doc.tables):
                rows = []
                for row in tbl.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    df = pd.DataFrame(rows)
                    df = df.dropna(how="all")
                    result[f"Table_{t_idx + 1}"] = df

            # Extract paragraph text as a single-column sheet (skip empties)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if paras:
                df_paras = pd.DataFrame(paras, columns=None)
                result["Text"] = df_paras

            if not result:
                # empty document
                result["Sheet1"] = pd.DataFrame()
            return result

        # .doc with no converter — read as plain text
        try:
            lines = open(path, "r", encoding="utf-8", errors="replace").read().splitlines()
            lines = [l for l in lines if l.strip()]
            df = pd.DataFrame(lines, columns=None)
            return {"Sheet1": df}
        except Exception:
            return {"Sheet1": pd.DataFrame()}

    # ── PDF — extract tables via pdfplumber, fall back to text ───────────
    if ext == ".pdf":
        try:
            import pdfplumber
            sheets: dict[str, pd.DataFrame] = {}
            with pdfplumber.open(path) as pdf:
                for p_idx, page in enumerate(pdf.pages):
                    for t_idx, table in enumerate(page.extract_tables() or []):
                        if table:
                            df = pd.DataFrame(table).dropna(how="all")
                            sheets[f"Page{p_idx+1}_Table{t_idx+1}"] = df
                if not sheets:
                    # No tables found — extract raw text as single-column sheet
                    lines = []
                    for page in pdf.pages:
                        txt = page.extract_text() or ""
                        lines.extend([l for l in txt.splitlines() if l.strip()])
                    sheets["Text"] = pd.DataFrame(lines, columns=None)
            return sheets
        except ImportError:
            # pdfplumber not available — use pypdf text extraction
            import pypdf
            reader = pypdf.PdfReader(path)
            lines = []
            for page in reader.pages:
                txt = page.extract_text() or ""
                lines.extend([l for l in txt.splitlines() if l.strip()])
            df = pd.DataFrame(lines, columns=None)
            return {"Text": df}

    # ── Unknown extension — try Excel first, then TSV ─────────────────────
    try:
        return _read_excel_sheets(path)
    except Exception:
        pass
    try:
        df = pd.read_csv(path, sep="	", header=None, dtype=str,
                         encoding_errors="replace", on_bad_lines="skip")
        df = df.dropna(how="all")
        return {"Sheet1": df}
    except Exception as e:
        raise ValueError(f"Unsupported file format for: {path}") from e


def _sheet_header_tokens(df: pd.DataFrame) -> list[str]:
    """Collect lowercase tokens from the first 3 rows (kept for backward compat)."""
    tokens: list[str] = []
    for _, row in df.head(3).iterrows():
        for val in row:
            if pd.notna(val):
                tokens.append(str(val).lower().strip())
    return tokens


def _classify_sheet(df: pd.DataFrame) -> ClassificationResult:
    """
    Classify a sheet using the spec-driven matcher (spec_match.py).
    Returns a ClassificationResult with confidence score and column gap report.
    """
    return classify_sheet(df)


def _count_data_rows(df: pd.DataFrame) -> int:
    """Approximate number of data rows (skip header rows)."""
    return max(0, len(df) - 2)


# ─────────────────────────────────────────────────────────────
# LLM metadata extraction
# ─────────────────────────────────────────────────────────────

def _extract_metadata_regex(pdf_text: str) -> dict:
    """
    Best-effort metadata extraction from raw PDF text using regex patterns.
    Used as a fallback when the LLM call fails or returns incomplete data.
    Tuned against Nature Communications / high-impact journal PDF structure.
    """
    import re as _re

    def _first(patterns, text, default=""):
        for pat in patterns:
            m = _re.search(pat, text, _re.IGNORECASE)
            if m:
                return m.group(1).strip()
        return default

    def _find_int(patterns, text, default="?"):
        for pat in patterns:
            m = _re.search(pat, text, _re.IGNORECASE)
            if m:
                for g in m.groups():
                    if g and _re.search(r"\d", g):
                        return g.strip()
        return default

    # ── Title: join lines after DOI line until author line starts ─────────
    doi_match  = _re.search(r"10\.[0-9]{4,}/\S+", pdf_text)
    title = "Study Title Not Detected"
    if doi_match:
        after_doi = pdf_text[doi_match.end():]
        title_lines = []
        for l in after_doi.splitlines():
            l = l.strip()
            if not l:
                continue
            # Author lines contain superscript digits attached to surnames
            if _re.search(r"[A-Z][a-z]+\d", l) or _re.search(r"\b[A-Z][a-z]+ [A-Z][a-z]+\d", l):
                break
            if len(l) > 5:
                title_lines.append(l)
            if len(title_lines) >= 3:
                break
        if title_lines:
            title = " ".join(title_lines)
    if title == "Study Title Not Detected":
        for l in pdf_text.splitlines():
            l = l.strip()
            if 20 < len(l) < 200 and any(w in l.lower() for w in
               ["genomic","transcriptom","landscape","characteriz","sequenc","mutati","cancer","tumor","tumour"]):
                title = l
                break

    # ── DOI ───────────────────────────────────────────────────────────────
    doi = _first([
        r"https?://doi\.org/([^\s,;)]+)",
        r"(?:doi|DOI)[:\s]+([10]\.[0-9]{4,}/\S+)",
        r"\b(10\.[0-9]{4,}/[^\s,;)\]]+)",
    ], pdf_text).rstrip(".,;)")

    # ── PMID ──────────────────────────────────────────────────────────────
    pmid = _first([
        r"PMID[:\s]+(\d{6,9})",
        r"pubmed\.ncbi\.nlm\.nih\.gov/(\d{6,9})",
    ], pdf_text)

    # ── Year: prefer Accepted > Published > Received ──────────────────────
    year = (_first([r"Accepted[:\s]+\d+\s+\w+\s+(\d{4})"], pdf_text) or
            _first([r"Published[:\s]+\d+\s+\w+\s+(\d{4})",
                    r"Published online[:\s]+(\d{4})"], pdf_text) or
            _first([r"Received[:\s]+\d+\s+\w+\s+(\d{4})"], pdf_text) or
            _first([r"\b(20[12][0-9])\b"], pdf_text))

    # ── Journal ───────────────────────────────────────────────────────────
    journal = _first([
        r"(Nature Communications)",
        r"(Nature\s+(?:Genetics|Medicine|Cancer|Methods|Biotechnology|Chemical Biology|Cell Biology|Immunology))",
        r"(Nature\s+\w+)",
        r"(Cancer\s+(?:Cell|Discovery|Research|Medicine))",
        r"(Cell\s+(?:Genomics|Systems|Reports|Research|Host|Stem))",
        r"(Science\s+(?:Translational|Advances|Medicine))",
        r"(Clinical\s+Cancer\s+Research)",
        r"(Journal\s+of\s+Clinical\s+Oncology)",
        r"(Genome\s+(?:Research|Biology|Medicine))",
        r"(Blood\s+Cancer\s+Journal)",
        r"(Leukemia)",
        r"(\bNEJM\b|New England Journal)",
    ], pdf_text)

    # ── First author surname ─────────────────────────────────────────────
    # Strategy: find author list line after title (right after DOI), extract
    # the SURNAME of the first author.  Author lines look like:
    #   "Feifei Xie1,10,S h u z h e n..."  or  "Smith J1, Jones A2, ..."
    author = ""
    if doi_match:
        after_doi = pdf_text[doi_match.end():]
        for l in after_doi.splitlines():
            l = l.strip()
            if not l:
                continue
            # Pattern 1: "Firstname Surname[digit]" — e.g. "Feifei Xie1,10,"
            am = _re.match(r"[A-Z][a-z]+\s+([A-Z][a-z]+)\d", l)
            if am:
                author = am.group(1)
                break
            # Pattern 2: "Surname, Initials" style — "Smith J1,"
            am2 = _re.match(r"([A-Z][a-z]{1,15}),\s*[A-Z]\.?\s*\d", l)
            if am2:
                author = am2.group(1)
                break
            # Stop if we've passed the author block (abstract/intro begins)
            if len(l) > 150 or _re.match(r"[A-Z][a-z].{80,}", l):
                break
    # Fallback: "Surname et al." anywhere in text
    if not author:
        m = _re.search(r"([A-Z][a-z]{2,15})\s+et\s+al", pdf_text[:3000])
        if m:
            author = m.group(1)

    # ── Reference genome ──────────────────────────────────────────────────
    genome = _first([
        r"\b(hg38|GRCh38)\b",
        r"\b(hg19|GRCh37)\b",
        r"aligned\s+to\s+(hg\d+|GRCh\d+)",
        r"reference\s+genome[:\s]+(hg\d+|GRCh\d+)",
        r"mapped\s+to\s+(hg\d+|GRCh\d+)",
        r"NCBI\s+[Bb]uild\s+(37|38)",
    ], pdf_text)
    if genome:
        genome = (genome
                  .replace("GRCh38","hg38").replace("GRCh37","hg19")
                  .replace("38","hg38").replace("37","hg19"))
        # Normalise aliases
        if genome not in ("hg19","hg38"):
            genome = "hg38"
    else:
        # Cannot determine from text — leave blank so the LLM or user can fill in
        genome = ""

    # ── Sample / patient counts ───────────────────────────────────────────
    n_samples = _find_int([
        r"(\d+)\s+(?:tumor|tumour|cancer|primary)\s+samples?",
        r"(\d+)\s+GISTs?\b",
        r"(\d+)\s+samples?\s+(?:were|from|across|in|with)",
        r"(?:total\s+of\s+)?(\d{2,4})\s+samples?",
        r"n\s*=\s*(\d+)\s+samples?",
        r"(\d+)\s+(?:tumor|tumour)\s+(?:specimens?|biopsies|cases)",
    ], pdf_text)
    n_patients = _find_int([
        r"(\d+)\s+patients?",
        r"(\d+)\s+(?:individuals?|subjects?|donors?)",
        r"cohort\s+of\s+(\d+)",
        r"(\d+)\s+cases?\b",
    ], pdf_text)

    # ── Sequencing types ──────────────────────────────────────────────────
    seq_types = []
    for label, patterns in [
        ("WES",       [r"\bWES\b", r"whole[- ]exome\s+seq"]),
        ("WGS",       [r"\bWGS\b", r"whole[- ]genome\s+seq"]),
        ("WTS",       [r"\bWTS\b", r"whole[- ]transcriptome\s+seq"]),
        ("RNA-seq",   [r"\bRNA-?seq\b"]),
        ("scRNA-seq", [r"\bscRNA-?seq\b", r"single[- ]cell\s+RNA"]),
        ("ChIP-seq",  [r"\bChIP-?seq\b"]),
        ("ATAC-seq",  [r"\bATAC-?seq\b"]),
        ("targeted",  [r"targeted\s+(?:sequencing|panel|NGS)"]),
    ]:
        if any(_re.search(p, pdf_text, _re.IGNORECASE) for p in patterns):
            seq_types.append(label)

    # ── Cancer type ───────────────────────────────────────────────────────
    cancer_map = [
        (r"\bGIST\b|gastrointestinal\s+stromal",     "gist",  "Gastrointestinal Stromal Tumor"),
        (r"\bbreast\s+cancer\b|\bBRCA\b",           "brca",  "Breast Invasive Carcinoma"),
        (r"\blung\s+adenocarcinoma\b|\bLUAD\b",     "luad",  "Lung Adenocarcinoma"),
        (r"\blung\s+squamous\b|\bLUSC\b",           "lusc",  "Lung Squamous Cell Carcinoma"),
        (r"\bnon-small\s+cell\s+lung\b|\bNSCLC\b", "nsclc", "Non-Small Cell Lung Cancer"),
        (r"\blung\s+cancer\b",                        "luad",  "Lung Cancer"),
        (r"\bcolorectal\b|\bCRC\b|\bCOAD\b",       "coad",  "Colorectal Adenocarcinoma"),
        (r"\bglioblastoma\b|\bGBM\b",                "gbm",   "Glioblastoma Multiforme"),
        (r"\bglioma\b",                                "lgggbm","Glioma"),
        (r"\bmelanoma\b|\bSKCM\b",                   "skcm",  "Skin Cutaneous Melanoma"),
        (r"\bpancreatic\s+(?:cancer|ductal|adenocarcinoma)\b|\bPAAD\b", "paad", "Pancreatic Adenocarcinoma"),
        (r"\bprostate\s+cancer\b|\bPRAD\b",         "prad",  "Prostate Adenocarcinoma"),
        (r"\bovarian\s+(?:cancer|carcinoma)\b|\bOV\b", "ov", "Ovarian Serous Cystadenocarcinoma"),
        (r"\bhepat\w+\s+(?:carcinoma|cancer)\b|\bHCC\b", "hcc", "Hepatocellular Carcinoma"),
        (r"\bgastric\s+(?:cancer|carcinoma)\b|\bSTAD\b", "stad", "Stomach Adenocarcinoma"),
        (r"\bladder\s+(?:cancer|carcinoma)\b|\bBLCA\b", "blca", "Bladder Urothelial Carcinoma"),
        (r"\bacute\s+myeloid\s+leukemia\b|\bAML\b", "aml", "Acute Myeloid Leukemia"),
        (r"\bchronic\s+lymphocytic\s+leukemia\b|\bCLL\b", "cll", "Chronic Lymphocytic Leukemia"),
        (r"\bleukemia\b",                              "leuk",  "Leukemia"),
        (r"\blymphoma\b|\bDLBCL\b",                  "dlbcl", "Diffuse Large B-Cell Lymphoma"),
        (r"\bmultiple\s+myeloma\b|\bMM\b",          "mm",    "Multiple Myeloma"),
        (r"\brenal\s+(?:cell\s+carcinoma|cancer)\b|\bRCC\b|\bKIRC\b", "kirc", "Renal Clear Cell Carcinoma"),
        (r"\bthyroid\s+(?:cancer|carcinoma)\b|\bTHCA\b", "thca", "Thyroid Carcinoma"),
        (r"\bendometrial\b|\buterine\b|\bUCEC\b",  "ucec",  "Uterine Corpus Endometrial Carcinoma"),
        (r"\bsarcoma\b",                               "sarc",  "Sarcoma"),
        (r"\bmesothelioma\b|\bMESO\b",               "meso",  "Mesothelioma"),
        (r"\bcervical\b|\bCESC\b",                   "cesc",  "Cervical Squamous Cell Carcinoma"),
        (r"\bhead\s+and\s+neck\b|\bHNSC\b",        "hnsc",  "Head and Neck Squamous Cell Carcinoma"),
    ]
    cancer_t, cancer_full = "mixed", "Mixed Cancer Type"
    for pat, ct, cf in cancer_map:
        if _re.search(pat, pdf_text[:3000], _re.IGNORECASE):
            cancer_t, cancer_full = ct, cf
            break

    study_id = f"{cancer_t}_{author.lower()}_{year}" if author and year else f"{cancer_t}_study_{year or '2024'}"
    study_id = _re.sub(r"[^a-z0-9_]", "_", study_id).strip("_")

    # ── Data repositories ─────────────────────────────────────────────────
    repos = []
    for pat in [r"(GSE\d{5,7})", r"(EGAS\d{11})", r"(phs\d{6,7})",
                r"(HRA\d{6})", r"(PRJNA\d+)", r"(SRP\d+)",
                r"(ERP\d+)", r"dbGaP\s+accession[:\s]+(\S+)"]:
        for m in _re.finditer(pat, pdf_text, _re.IGNORECASE):
            v = m.group(1).strip()
            if v not in repos:
                repos.append(v)

    # ── Corresponding author ──────────────────────────────────────────────
    corresp = _first([
        r"[Cc]orresponding\s+authors?[:\s]+([^\n]{10,100})",
        r"[Cc]orrespondence[:\s]+([^\n]{10,100})",
        r"\*?[Ee]-mail[:\s]+([^\n]{10,80})",
    ], pdf_text)

    # ── Key findings ──────────────────────────────────────────────────────
    # Try to extract sentences with result-indicating words
    key_findings = []
    for sent in _re.split(r"[.!?]\s+", pdf_text[:4000]):
        sent = sent.strip()
        if (len(sent) > 40 and
            any(w in sent.lower() for w in ["identified","revealed","found","discover",
               "demonstrate","show","report","novel","significant","recurrent"]) and
            len(key_findings) < 5):
            key_findings.append(sent[:150])

    # ── Build description ─────────────────────────────────────────────────
    seq_str = "/".join(seq_types[:3]) if seq_types else "genomic"
    n_str   = f" of {n_samples} samples" if n_samples != "?" else ""
    desc    = f"{seq_str.upper()} characterization{n_str} of {cancer_full}."
    if author and journal and year:
        desc = f"{seq_str} study{n_str} of {cancer_full}. Published in {journal} ({year})."

    return {
        "study_title":          title[:200],
        "cancer_type":          cancer_t,
        "cancer_type_full":     cancer_full,
        "num_samples":          n_samples,
        "num_patients":         n_patients,
        "reference_genome":     genome,
        "sequencing_types":     seq_types,
        "pmid":                 pmid,
        "doi":                  doi,
        "first_author_surname": author,
        "year":                 year,
        "journal":              journal,
        "study_id_suggestion":  study_id,
        "description":          desc,
        "meta_description":     desc[:200],
        "key_findings":         key_findings,
        "primary_site":         "",
        "cohort_description":   (f"{n_patients} patients, {n_samples} samples."
                                  if n_patients != "?" else ""),
        "data_repositories":    repos[:4],
        "corresponding_authors": corresp,
    }

def _extract_metadata_llm(pdf_text: str, model: str, temperature: float) -> dict:
    import json, logging
    llm = load_chat_model(model)
    messages = [
        SystemMessage(content=SYSTEM_PROMPT_CURATOR),
        HumanMessage(content=pdf_text[:8000]),
    ]
    response = llm.invoke(messages)
    raw = response.content.strip()
    raw = re.sub(r"^```[a-z]*\n?", "", raw)
    raw = re.sub(r"\n?```$",       "", raw)
    try:
        llm_data = json.loads(raw)
    except Exception as e:
        logging.warning(f"LLM JSON parse failed ({e}); using regex fallback.")
        llm_data = {}
    # Merge: regex fills in anything the LLM left blank/missing
    fallback = _extract_metadata_regex(pdf_text)
    merged   = {**fallback}   # start with regex values
    for key, val in llm_data.items():
        if val and val not in ("?", "...", "Unknown", "mixed", "study_2024", ""):
            merged[key] = val
    return merged


# ─────────────────────────────────────────────────────────────
# Per-format curation instructions
# ─────────────────────────────────────────────────────────────

def _build_instructions(cr: ClassificationResult, df: pd.DataFrame, sheet_name: str) -> dict:
    """Return structured curation instructions for a classified sheet."""
    fmt = cr.format_key
    cols = [str(c) for c in df.iloc[0] if pd.notna(c)]  # row 0 is usually the real header

    base = {
        "classification":      fmt,
        "cbio_target_file":    cr.target_file,
        "curability":          CURABILITY.get(fmt, ("NO", "N/A"))[0],
        "priority":            CURABILITY.get(fmt, ("NO", "N/A"))[1],
        "confidence":          cr.confidence,
        "verdict":             cr.verdict,
        "spec_source":         cr.spec_source,
        "spec_fetched_at":     cr.spec_fetched_at,
        "detected_columns":    cols[:20],
        "required_present":    cr.required_present,
        "required_missing":    cr.required_missing,
        "optional_present":    cr.optional_present,
        "alias_mappings":      cr.detected_as_aliases,
        "candidate_scores":    cr.all_scores,
        "data_rows_approx":    _count_data_rows(df),
        "column_mapping":      [],
        "transformations":     [],
        "missing_required":    [],
        "notes":               [cr.notes] if cr.notes else [],
    }

    # ── CLINICAL_PATIENT ──────────────────────────────────────
    if fmt == "CLINICAL_PATIENT":
        base["column_mapping"] = [
            ("Case ID / Patient ID", "PATIENT_ID", "Required. Must be unique per patient."),
            ("Gender / Sex",         "SEX",         "STRING: Male / Female"),
            ("Age",                  "AGE",          "NUMBER: age at diagnosis in years"),
            ("Primary Location",     "TUMOR_SITE",   "STRING: anatomical site"),
        ]
        base["missing_required"] = [
            "OS_STATUS  (0:LIVING / 1:DECEASED) — needed for survival plots",
            "OS_MONTHS  — overall survival in months since diagnosis",
        ]
        base["transformations"] = [
            "Add 5-row cBioPortal header (Display Name, Description, Datatype, Priority, Column Name)",
            "Separate patient rows from sample rows if combined in one sheet",
        ]

    # ── CLINICAL_SAMPLE ───────────────────────────────────────
    elif fmt == "CLINICAL_SAMPLE":
        base["column_mapping"] = [
            ("Sample ID",      "SAMPLE_ID",   "Required. Must match Tumor_Sample_Barcode in MAF."),
            ("Case ID",        "PATIENT_ID",  "Required. Link to patient file."),
            ("Sample Type",    "SAMPLE_TYPE", "STRING: Primary / Metastatic / Recurrence / etc."),
            ("Risk / Stage",   "RISK_STRATIFICATION", "STRING: Low-risk / High-risk / Metastatic"),
            ("TKI treatment",  "TKI_TREATMENT",       "STRING: TKI naive / TKI treated"),
        ]
        base["transformations"] = [
            "Add 5-row cBioPortal clinical header",
            "Tumour / normal sample rows: include only tumour sample rows",
            "Any ≤/≥ range values (e.g. tumour size '≤2 cm') — keep as STRING datatype",
        ]

    # ── MUTATION_MAF ──────────────────────────────────────────
    elif fmt == "MUTATION_MAF":
        base["column_mapping"] = [
            ("Gene / Gene Symbol",               "Hugo_Symbol",            "HGNC-validated gene name"),
            ("Sample ID",                        "Tumor_Sample_Barcode",   "Must match SAMPLE_ID in clinical file"),
            ("Chromosome",                       "Chromosome",             "Strip 'chr' prefix → integer 1-22, X, Y"),
            ("Start Location",                   "Start_Position",         "Integer (1-based)"),
            ("End Location",                     "End_Position",           "Integer (1-based)"),
            ("Reference",                        "Reference_Allele",       "ACGT string; '-' for insertions"),
            ("Alteration / Alt",                 "Tumor_Seq_Allele2",      "ACGT string; '-' for deletions"),
            ("Exonic Mutation Function",         "Variant_Classification", "See ANNOVAR→MAF mapping below"),
            ("Mutation Type (exonic/splicing)",  "Variant_Type",           "SNP / INS / DEL"),
            ("VAF / Variant Allele Frequency",   "t_AF (custom)",          "Float 0–1; t_depth optional"),
        ]
        base["transformations"] = [
            "ANNOVAR → MAF Variant_Classification: " + 
            " | ".join(f"{k} → {v}" for k, v in ANNOVAR_TO_MAF.items()),
            "Strip 'chr' prefix from Chromosome column",
            "Add NCBI_Build = GRCh37 (if hg19 reference used)",
            "Add Matched_Norm_Sample_Barcode from matched normal sample IDs",
            "Set Strand = + for all entries (default)",
        ]
        base["missing_required"] = [
            "t_depth, t_ref_count, t_alt_count, n_ref_count, n_alt_count (tumour/normal depths)",
            "Matched_Norm_Sample_Barcode (derive from clinical file: e.g., '1T' → '1N')",
            "NCBI_Build (set to GRCh37 for hg19)",
        ]

    # ── STRUCTURAL_VARIANT ────────────────────────────────────
    elif fmt == "STRUCTURAL_VARIANT":
        base["column_mapping"] = [
            ("Sample ID",          "Sample_Id",         "Must match SAMPLE_ID in clinical file"),
            ("Left Gene / Gene 1", "Site1_Hugo_Symbol", "First gene at breakpoint"),
            ("Right Gene / Gene 2","Site2_Hugo_Symbol", "Second gene at breakpoint (or same gene for intragenic)"),
            ("Start Chr / Left Chr","Site1_Chromosome", "Strip 'chr' prefix"),
            ("End Chr / Right Chr", "Site2_Chromosome", "Strip 'chr' prefix"),
            ("Start Position",     "Site1_Position",    "Integer coordinate"),
            ("End Position",       "Site2_Position",    "Integer coordinate"),
            ("SV Type",            "Class",             "DELETION / DUPLICATION / INVERSION / TRANSLOCATION / FUSION"),
        ]
        base["transformations"] = [
            "Remap SV Type: DEL→DELETION, DUP→DUPLICATION, INV→INVERSION, BND→TRANSLOCATION",
            "For fusion files: set Class = FUSION",
            "Add SV_Status = SOMATIC for all entries",
            "For multi-caller fusion files: select consensus breakpoint (≥2 callers agree)",
            "Split multi-gene 'Gene Involved' strings on '/' to populate Site1/Site2",
        ]
        base["missing_required"] = [
            "SV_Status (set to SOMATIC)",
            "Event_Info (description of the event)",
            "Site1_Region / Site2_Region (EXONIC / INTRONIC / INTERGENIC)",
        ]

    # ── DISCRETE_CNA / GISTIC ─────────────────────────────────
    elif fmt == "DISCRETE_CNA":
        base["transformations"] = [
            "Separate AMP and DEL rows into data_gistic_genes_amp.txt and data_gistic_genes_del.txt",
            "Gene list in 'All Genes in Wide Peak' → parse into per-gene rows",
            "For per-sample discrete matrix: GISTIC calls → integer matrix (-2/-1/0/1/2)",
            "Request full GISTIC output directory from authors for complete loading",
        ]
        base["missing_required"] = [
            "Per-sample gene-level CNA matrix (data_CNA.txt) — not in summary tables",
            "GISTIC all_lesions.conf_90 file for arm-level per-sample calls",
        ]

    # ── MUTSIG ────────────────────────────────────────────────
    elif fmt == "MUTSIG":
        base["column_mapping"] = [
            ("gene",    "gene",    "Hugo Symbol"),
            ("codelen", "codelen", "Coding length in bp"),
            ("nnon",    "nnon",    "Non-silent mutation count"),
            ("npat",    "npat",    "Number of patients with mutation"),
        ]
        base["missing_required"] = [
            "p  — p-value (retrieve from full MutSigCV output)",
            "q  — q-value / FDR (retrieve from full MutSigCV output)",
        ]

    # ── GENERIC_ASSAY ─────────────────────────────────────────
    elif fmt == "GENERIC_ASSAY":
        base["transformations"] = [
            "Transpose to entity × sample matrix format required by Generic Assay",
            "Define GENERIC_ASSAY_TYPE in meta file (e.g. MUTATIONAL_SIGNATURE, ARM_LEVEL_CNA)",
            "Add ENTITY_STABLE_ID column as first column",
            "Per-sample values must be numeric",
        ]
        base["notes"] = [
            "Immune deconvolution scores (CIBERSORT, ESTIMATE) → Generic Assay: IMMUNE_SCORE",
            "Mutational signature exposures → Generic Assay: MUTATIONAL_SIGNATURE",
            "Arm-level CNA frequencies → Generic Assay: ARM_LEVEL_CNA (needs per-sample calls)",
        ]

    # ── METHYLATION ───────────────────────────────────────────
    elif fmt == "METHYLATION":
        base["transformations"] = [
            "Convert raw clone counts to beta values (methylated clones / total clones)",
            "Aggregate CpG sites to gene-level (mean or median beta value per gene)",
            "Format: gene × sample matrix with float beta values (0–1)",
        ]
        base["missing_required"] = [
            "Gene-level aggregation (current data is per-CpG-site)",
            "Broader sample coverage (only a few samples in supplement)",
        ]

    # ── Build narrative intro paragraph for the report ──────────────────
    fmt = base["classification"]
    n_rows = base.get("data_rows_approx", 0)
    tgt    = base.get("cbio_target_file", "")
    verd   = base.get("verdict", "")
    n_cols = len(base.get("detected_columns", []))

    if fmt == "CLINICAL_PATIENT":
        base["intro"] = (
            f"{verd}  These should be split into a patient-level and a "
            "sample-level clinical file using the column mapping tables below."
        ).strip()
    elif fmt == "CLINICAL_SAMPLE":
        base["intro"] = (
            f"{verd}  Split into sample-level attributes and merge with the "
            "patient file for complete clinical data."
        ).strip()
    elif fmt == "MUTATION_MAF":
        base["intro"] = (
            f"{verd}  The data is ANNOVAR-annotated and largely maps to MAF "
            "format, but column remapping is required."
        ).strip()
    elif fmt == "STRUCTURAL_VARIANT":
        base["intro"] = (
            f"{verd}  Both DNA SVs and RNA fusions map to cBioPortal's data_sv.txt format."
        ).strip()
    elif fmt in ("DISCRETE_CNA", "CONTINUOUS_CNA", "GISTIC"):
        base["intro"] = (
            f"{verd}  For full cBioPortal GISTIC loading, the complete GISTIC output "
            "directory is needed (not always available in supplementary). "
            "However, partial curation is possible."
        ).strip()
    elif fmt == "MUTSIG":
        base["intro"] = (
            f"{verd}  cBioPortal's MutSig format requires gene, rank, p-value, and q-value."
        ).strip()
    else:
        base["intro"] = verd

    return base


# ─────────────────────────────────────────────────────────────
# Main orchestration
# ─────────────────────────────────────────────────────────────

def _analyse_supplementary_files(supp_paths: list[str]) -> list[dict]:
    """Inspect every sheet in every supp file and return analysis records.

    Supports: .xlsx, .xls, .csv, .tsv, .tab, .txt, .maf, .doc, .docx
    """
    records: list[dict] = []
    for path in supp_paths:
        fname = Path(path).name
        try:
            sheets = _read_file_as_sheets(path)
        except Exception as e:
            records.append({
                "file": fname, "sheet": "-",
                "classification": "NOT_LOADABLE",
                "error": str(e),
                "cbio_target_file": "N/A",
                "curability": "NO", "priority": "N/A",
                "confidence": 0, "verdict": f"Parse error: {e}",
                "required_present": [], "required_missing": [],
                "optional_present": [], "alias_mappings": {},
                "candidate_scores": [],
                "detected_columns": [], "data_rows_approx": 0,
                "column_mapping": [], "transformations": [],
                "missing_required": [], "notes": [],
            })
            continue

        for sheet_name, df in sheets.items():
            cr = _classify_sheet(df)
            instructions = _build_instructions(cr, df, sheet_name)
            instructions["file"] = fname
            instructions["sheet"] = sheet_name
            records.append(instructions)

    return records


# ─────────────────────────────────────────────────────────────
# DOCX report builder
# ─────────────────────────────────────────────────────────────

# Colour palette
C_DARK  = "1F4E79"
C_MID   = "2C5F8A"
C_LIGHT = "2E75B6"
C_GREEN = "375623"  # fill E2EFDA
C_AMBER = "7F6000"  # fill FFF2CC
C_RED   = "843C0C"  # fill FCE4D6
C_GREY  = "595959"

FILL_GREEN = "E2EFDA"
FILL_AMBER = "FFF2CC"
FILL_RED   = "FCE4D6"
FILL_GREY  = "F2F2F2"
FILL_HEAD  = "2C5F8A"
FILL_ALT   = "F7FAFD"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, texts: list[str], widths_cm: list[float]):
    row = table.add_row()
    for i, (text, w) in enumerate(zip(texts, widths_cm)):
        cell = row.cells[i]
        cell.width = Inches(w / 2.54)
        _set_cell_bg(cell, FILL_HEAD)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _data_row(table, texts: list[str], widths_cm: list[float],
              alt: bool = False, status_col: int | None = None):
    row = table.add_row()
    fill_map = {
        "YES":     (FILL_GREEN, C_GREEN),
        "PARTIAL": (FILL_AMBER, C_AMBER),
        "NO":      (FILL_RED,   C_RED),
        "HIGH":    (FILL_RED,   C_RED),
        "MEDIUM":  (FILL_AMBER, C_AMBER),
        "LOW":     (FILL_GREY,  C_GREY),
        "N/A":     (FILL_GREY,  C_GREY),
    }
    bg = FILL_ALT if alt else "FFFFFF"
    for i, (text, w) in enumerate(zip(texts, widths_cm)):
        cell = row.cells[i]
        cell.width = Inches(w / 2.54)
        if status_col is not None and i == status_col and text in fill_map:
            fill, col = fill_map[text]
            _set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(col)
        else:
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)


def _add_heading(doc: DocxDocument, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(C_DARK)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), C_MID)
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(C_MID)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(C_LIGHT)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def _add_para(doc: DocxDocument, text: str, bold_prefix: str = "",
              italic: bool = False, color: str = "000000"):
    p = doc.add_paragraph()
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_bullet(doc: DocxDocument, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p




def _build_report(
    meta: dict,
    records: list[dict],
    output_path: str,
):
    """
    Generate a cBioPortal curation report matching the v2 reference format exactly.

    Exact document blueprint (from cbioportal_curation_report_GIST_Xie2024_v2.docx):
    ─────────────────────────────────────────────────────────────────────────────────
    P000  Title          "cBioPortal Curation Report"
    P001  Subtitle       study title
    P002  Citation       "Author et al., Journal (year) vol  |  DOI: ...  |  PMID: ..."
    P003  H1             "1. Study Overview"
    P004  body           "Study:  ..."
    P005  body           "Cohort:  ..."
    P006  body           "Reference Genome:  ..."
    P007  body           "Primary Site:  ... Cancer Type: ..."
    P008  body           "Key Findings:  ..."
    P009  empty
    P010  H1             "2. Proposed cBioPortal Study Metadata (meta_study.txt)"
    P011  empty
    TABLE 0 (9r x 2c)   meta_study.txt fields
    P012  empty
    P013  H1             "3. Supplementary File Inventory & Curability Assessment"
    P014  body           "Summary of all N supplementary data files..."
    P015  empty
    TABLE 1 (N x 5c)    file inventory
    P016  empty
    P017  H1             "4. Detailed Curation Instructions by File Type"
    P018  H2             "4.1  Clinical Data ..."
    P019  body           clinical intro
    P020  empty
    P021  H3             "Patient-level attributes ..."
    P022  empty
    TABLE 2 (N x 4c)    patient column mapping
    P023  empty
    P024  H3             "Sample-level attributes ..."
    P025  empty
    TABLE 3 (N x 4c)    sample column mapping
    P026  empty
    P027  body           "File header format reminder: ..."
    P028  empty
    P029  H2             "4.2  Mutation Data ..."
    P030  body           MAF intro
    P031  empty
    TABLE 4 (N x 3c)    MAF column mapping
    P032  empty
    P033  body           "ANNOVAR -> MAF mapping: ..."
    P034  empty
    P035  H2             "4.3  Structural Variant Data ..."
    P036  body           SV intro
    P037  empty
    TABLE 5 (N x 3c)    SV column mapping
    P038  empty
    P039  note           "Note: For fusions ..."
    P040  empty
    P041  H2             "4.4  Copy Number Data ..."
    P042  body           CNA intro
    P043  empty
    P044  H3             "From Supp Data X -> data_gistic_genes_*.txt"
    P045+ bullets
    P_n   empty
    P_n1  H3             "From Supp Data X -> Generic Assay: Arm-Level CNA"
    P_n2+ bullets
    P_n3  empty
    P055  H2             "4.5  MutSig Data ..."
    P056  body           MutSig intro
    P057  empty
    TABLE 6 (N x 3c)    MutSig column mapping
    P058  empty
    P059  H2             "4.6  Gene Fusions ..."
    P060  body           fusion intro
    P061+ bullets
    P065  empty
    P066  H1             "5. Case Lists"
    P067  body           intro
    P068  empty
    TABLE 7 (N x 3c)    case lists
    P069  empty
    P070  H1             "6. Data Gaps & Contact with Authors"
    P071  body           intro
    P072  empty
    TABLE 8 (N x 2c)    data gaps
    P073  empty
    P074  H1             "7. Curation Checklist & Priority Order"
    P075  empty
    TABLE 9 (N x 4c)    checklist
    P076  empty
    P077  H1             "8. Data Accession & Reproducibility"
    P078+ body           "Label: value"  (one per repository / author line)
    P_n   empty
    P_n1  note           "All curation should follow..."
    """

    from docx import Document as DocxDocument
    from docx.shared import Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Exact colours from reference ───────────────────────────────────────
    C_DARK    = "1F4E79"   # H1 text, title
    C_MID     = "2C5F8A"   # H2 text, H1 border, table header bg
    C_LIGHT   = "2E75B6"   # H3 text
    C_GREY    = "595959"   # italic note/citation text

    C_HDR_BG  = "2C5F8A"   # table header row bg
    C_ODD     = "FFFFFF"   # data row odd
    C_EVEN    = "F0F5FB"   # data row even
    C_YES_BG  = "E2EFDA"; C_YES_TXT = "375623"
    C_PAR_BG  = "FFF2CC"; C_PAR_TXT = "7F6000"
    C_NA_BG   = "F2F2F2";  C_NA_TXT  = "595959"

    # ── Exact column widths (EMU measured from reference) ──────────────────
    W0 = [Emu(2032000), Emu(3911600)]
    W1 = [Emu(889000), Emu(1905000), Emu(1397000), Emu(508000), Emu(1244600)]
    W2 = [Emu(1524000), Emu(1270000), Emu(762000), Emu(2387600)]
    W3 = [Emu(1524000), Emu(1270000), Emu(762000), Emu(2387600)]
    W4 = [Emu(1778000), Emu(1778000), Emu(2387600)]
    W5 = [Emu(1524000), Emu(1524000), Emu(2895600)]
    W6 = [Emu(1524000), Emu(1270000), Emu(3149600)]
    W7 = [Emu(1270000), Emu(1905000), Emu(2768600)]
    W8 = [Emu(2286000), Emu(3657600)]
    W9 = [Emu(317500),  Emu(2857500), Emu(762000), Emu(2006600)]

    # ═══════════════════════════════════════════════════════════════════════
    # LOW-LEVEL HELPERS
    # ═══════════════════════════════════════════════════════════════════════

    def _shd(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _cell_margins(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcMar")):
            tcPr.remove(old)
        mar = OxmlElement("w:tcMar")
        for side in ("top", "bottom", "left", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"),    "80" if side in ("top","bottom") else "120")
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def _hdr_row(table, headers, widths):
        row = table.add_row()
        for cell, h, w in zip(row.cells, headers, widths):
            cell.width = w
            _shd(cell, C_HDR_BG)
            _cell_margins(cell)
            r = cell.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def _dat_row(table, values, widths, row_idx=0, status_col=None):
        bg = C_EVEN if row_idx % 2 == 1 else C_ODD
        row = table.add_row()
        for i, (cell, val, w) in enumerate(zip(row.cells, values, widths)):
            cell.width = w
            _cell_margins(cell)
            txt = str(val) if val is not None else ""
            p = cell.paragraphs[0]
            if i == status_col:
                if txt == "YES":
                    _shd(cell, C_YES_BG); r = p.add_run(txt)
                    r.bold = True; r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor.from_string(C_YES_TXT)
                elif txt == "PARTIAL":
                    _shd(cell, C_PAR_BG); r = p.add_run(txt)
                    r.bold = True; r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor.from_string(C_PAR_TXT)
                else:
                    _shd(cell, C_NA_BG); r = p.add_run("N/A")
                    r.bold = True; r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor.from_string(C_NA_TXT)
            else:
                _shd(cell, bg)
                r = p.add_run(txt)
                r.font.size = Pt(9)

    def _newtable(n_cols):
        t = doc.add_table(rows=1, cols=n_cols)
        t.style = "Table Grid"
        t._tbl.remove(t.rows[0]._tr)
        return t

    # ── paragraph builders ─────────────────────────────────────────────────

    def _empty():
        doc.add_paragraph()

    def _title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Emu(50800)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(24)
        r.font.color.rgb = RGBColor.from_string(C_DARK)

    def _subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Emu(38100)
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor.from_string(C_MID)

    def _citation_line(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Emu(127000)
        r = p.add_run(text)
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(C_GREY)

    def _h1(text):
        p = doc.add_paragraph(style="Heading 1")
        p.clear()
        p.paragraph_format.space_before = Emu(177800)
        p.paragraph_format.space_after  = Emu(63500)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(14)
        r.font.color.rgb = RGBColor.from_string(C_DARK)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), C_MID)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _h2(text):
        p = doc.add_paragraph(style="Heading 2")
        p.clear()
        p.paragraph_format.space_before = Emu(127000)
        p.paragraph_format.space_after  = Emu(50800)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(C_MID)

    def _h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Emu(88900)
        p.paragraph_format.space_after  = Emu(38100)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(C_LIGHT)

    def _body(bold_prefix, normal_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Emu(25400)
        p.paragraph_format.space_after  = Emu(25400)
        if bold_prefix:
            rb = p.add_run(bold_prefix); rb.bold = True; rb.font.size = Pt(10)
        if normal_text:
            rn = p.add_run(normal_text); rn.font.size = Pt(10)

    def _note(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Emu(25400)
        p.paragraph_format.space_after  = Emu(25400)
        r = p.add_run(text)
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(C_GREY)

    def _bullet(text):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.space_before = Emu(25400)
        p.paragraph_format.space_after  = Emu(25400)
        r = p.add_run(text); r.font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════════
    # EXTRACT META
    # ═══════════════════════════════════════════════════════════════════════

    study_title  = meta.get("study_title",          "Untitled Study")
    author       = meta.get("first_author_surname",  "")
    journal      = meta.get("journal",               "")
    year         = meta.get("year",                  "")
    doi          = meta.get("doi",                   "")
    pmid         = meta.get("pmid",                  "")
    study_id     = meta.get("study_id_suggestion",   "study_2024")
    cancer_t     = meta.get("cancer_type",           "mixed")
    cancer_full  = meta.get("cancer_type_full",      cancer_t)
    ref_genome   = meta.get("reference_genome",      "hg38")
    seq_types    = meta.get("sequencing_types",      [])
    n_patients   = meta.get("num_patients",          "?")
    n_samples    = meta.get("num_samples",           "?")
    description  = meta.get("description",           "")
    meta_desc    = meta.get("meta_description",      "")
    cohort_desc  = meta.get("cohort_description",    "")
    primary_site = meta.get("primary_site",          "")
    key_findings = meta.get("key_findings",          [])
    data_repos   = meta.get("data_repositories",     [])
    corresp      = meta.get("corresponding_authors", "")

    # ── Classify records ───────────────────────────────────────────────────
    def _recs(*cls_keywords):
        out = []
        for r in records:
            c = r.get("classification", "").upper()
            for kw in cls_keywords:
                if kw.upper() in c:
                    out.append(r); break
        return out

    def _files(*cls_keywords):
        return list(dict.fromkeys(
            r.get("file", "") for r in _recs(*cls_keywords) if r.get("file")
        ))

    def _files_compact(*cls_keywords):
        """Like _files but returns a compact joined string: 'Supp Data 1 + 5'
        by stripping repeated prefixes."""
        files = _files(*cls_keywords)
        if not files:
            return "Supp Data"
        # Strip common prefix for compactness
        # e.g. ["Supp Data 1", "Supp Data 5"] -> "Supp Data 1 + 5"
        prefix = ""
        for f in files:
            # Common prefix pattern: "Supp Data "
            if f.lower().startswith("supp data ") or f.lower().startswith("supplementary data "):
                idx = f.lower().find(" data ") + len(" data ")
                prefix = f[:idx]
                break
        if prefix and all(f.startswith(prefix) for f in files):
            return prefix + " + ".join(f[len(prefix):] for f in files)
        return " + ".join(files)

    clin_pat  = _recs("CLINICAL_PATIENT")
    clin_sam  = _recs("CLINICAL_SAMPLE")
    maf_recs  = _recs("MUTATION_MAF", "MUTATION")
    sv_recs   = _recs("STRUCTURAL_VARIANT")
    cna_recs  = _recs("DISCRETE_CNA", "CONTINUOUS_CNA", "GISTIC")
    ms_recs   = _recs("MUTSIG")
    fus_recs  = [r for r in records
                 if "fusion" in r.get("cbio_target_file", "").lower()
                 or "fusion" in r.get("classification", "").lower()]

    def _first_verdict(recs, fallback=""):
        for r in recs:
            v = r.get("verdict", "")
            if v: return v
        return fallback

    def _first_intro(recs, fallback=""):
        """Return a rich narrative intro: prefer 'intro' field, then 'description',
        then a 'verdict' that looks like a proper sentence (contains a period or comma)."""
        for r in recs:
            for key in ("intro", "description", "section_intro"):
                v = r.get(key, "")
                if v and len(v) > 50: return v
        # fallback to verdict only if it looks sentence-like
        for r in recs:
            v = r.get("verdict", "")
            if v and ("." in v or "," in v) and len(v) > 50: return v
        return fallback

    def _first_notes(recs, fallback=""):
        for r in recs:
            notes = r.get("notes") or []
            if notes: return str(notes[0])
        return fallback

    def _first_transforms(recs):
        for r in recs:
            t = r.get("transformations") or []
            if t: return [str(x) for x in t]
        return []

    def _first_colmap(recs):
        for r in recs:
            cm = r.get("column_mapping") or []
            if cm: return cm
        return []

    # ── Build citation line ─────────────────────────────────────────────────
    # "Xie et al., Nature Communications (2024) 15:9495  |  DOI: ...  |  PMID: ..."
    # journal field may already contain "(year) vol:issue", e.g. "Nature Communications (2024) 15:9495"
    cite_parts = []
    if author:   cite_parts.append(f"{author} et al.")
    if journal:  cite_parts.append(journal)
    elif year:   cite_parts.append(f"({year})")
    citation_str = ", ".join(cite_parts)
    if doi:  citation_str += f"  |  DOI: {doi}"
    if pmid: citation_str += f"  |  PMID: {pmid}"

    # Short journal name for meta_study.txt citation row
    journal_name = journal.split("(")[0].strip().split(",")[0].strip() if journal else ""
    cite_short   = (f"{author} et al., {journal_name} {year}".strip(" ,") if author
                    else f"{journal_name} {year}".strip())

    # ═══════════════════════════════════════════════════════════════════════
    # BUILD DOCUMENT
    # ═══════════════════════════════════════════════════════════════════════
    doc = DocxDocument()
    for sec in doc.sections:
        sec.page_width    = Emu(7772400)
        sec.page_height   = Emu(10058400)
        sec.left_margin   = sec.right_margin = Emu(914400)
        sec.top_margin    = sec.bottom_margin = Emu(914400)

    # ── P000  Title ────────────────────────────────────────────────────────
    _title("cBioPortal Curation Report")

    # ── P001  Subtitle ─────────────────────────────────────────────────────
    _subtitle(study_title)

    # ── P002  Citation ─────────────────────────────────────────────────────
    _citation_line(citation_str)

    # ──────────────────────────────────────────────────────────────────────
    # §1  STUDY OVERVIEW
    # ──────────────────────────────────────────────────────────────────────
    _h1("1. Study Overview")   # P003

    study_line = description or (
        f"Multi-omics study of {n_samples} samples from {n_patients} patients. "
        + (", ".join(str(s) for s in seq_types) + ". " if seq_types else "")
        + (f"Published in {journal_name}, {year}." if journal_name and year else "")
    ).strip()
    _body("Study:  ", study_line)   # P004

    cohort_line = cohort_desc or f"{n_patients} patients, {n_samples} samples."
    _body("Cohort:  ", cohort_line)   # P005

    _body("Reference Genome:  ",   # P006
          f"{ref_genome} (UCSC). All coordinates in supplementary files reference {ref_genome}.")

    site_text = (f"{primary_site}. " if primary_site else "") + \
                f"Cancer Type: {cancer_full} ({cancer_t})"
    _body("Primary Site:  ", site_text)   # P007

    findings_text = "  |  ".join(str(f) for f in key_findings) if key_findings \
                    else "See paper for key genomic findings."
    _body("Key Findings:  ", findings_text)   # P008

    _empty()   # P009

    # ──────────────────────────────────────────────────────────────────────
    # §2  STUDY METADATA
    # ──────────────────────────────────────────────────────────────────────
    _h1("2. Proposed cBioPortal Study Metadata (meta_study.txt)")   # P010
    _empty()   # P011

    # TABLE 0 — meta_study.txt (9r x 2c)
    desc_val = meta_desc or (description[:200] if description else "")
    study_name = (f"{cancer_full} ({author} et al., {year})" if author and year
                  else study_title)
    t0 = _newtable(2)
    _hdr_row(t0, ["Field", "Value"], W0)
    for i, (field, value) in enumerate([
        ("cancer_study_identifier", study_id),
        ("type_of_cancer",          cancer_t),
        ("name",                    study_name),
        ("description",             desc_val),
        ("citation",                cite_short),
        ("pmid",                    pmid),
        ("reference_genome",        ref_genome),
        ("add_global_case_list",    "true"),
    ]):
        _dat_row(t0, [field, str(value)], W0, row_idx=i)

    _empty()   # P012

    # ──────────────────────────────────────────────────────────────────────
    # §3  FILE INVENTORY
    # ──────────────────────────────────────────────────────────────────────
    _h1("3. Supplementary File Inventory & Curability Assessment")   # P013

    _body("", f"Summary of all {len(records)} supplementary data files assessed against "
              "cBioPortal file format specifications "
              "(https://docs.cbioportal.org/file-formats/).")   # P014
    _empty()   # P015

    # TABLE 1 — inventory (N x 5c)
    t1 = _newtable(5)
    _hdr_row(t1, ["File", "Content", "cBioPortal Format", "Curate?", "Notes"], W1)
    for i, rec in enumerate(records):
        cur = str(rec.get("curability", "NO")).upper()
        if cur not in ("YES", "PARTIAL"):
            cur = "N/A"
        verdict = rec.get("verdict", "")
        content = verdict[:55] if verdict else str(rec.get("sheet", ""))
        fmt     = str(rec.get("cbio_target_file", ""))
        notes   = rec.get("notes") or rec.get("transformations") or []
        missing = rec.get("missing_required") or []
        note_txt = (str(notes[0])[:70] if notes
                    else ("⚠ " + str(missing[0])[:65]) if missing else "")
        _dat_row(t1,
                 [rec.get("file",""), content, fmt, cur, note_txt],
                 W1, row_idx=i, status_col=3)

    _empty()   # P016

    # ──────────────────────────────────────────────────────────────────────
    # §4  DETAILED CURATION INSTRUCTIONS
    # ──────────────────────────────────────────────────────────────────────
    _h1("4. Detailed Curation Instructions by File Type")   # P017

    # ── 4.1  CLINICAL ─────────────────────────────────────────────────────
    clin_all_files = list(dict.fromkeys(
        _files("CLINICAL_PATIENT") + _files("CLINICAL_SAMPLE")
    ))
    clin_src = _files_compact("CLINICAL_PATIENT", "CLINICAL_SAMPLE")
    _h2(f"4.1  Clinical Data \u2014 {clin_src}  \u2192  "
        "data_clinical_patient.txt / data_clinical_sample.txt")   # P018

    clin_intro = (
        _first_intro(clin_pat) or _first_intro(clin_sam) or
        f"The primary clinical file contains {n_patients} patients and {n_samples} samples. "
        "These should be split into a patient-level and a sample-level clinical file."
    )
    _body("", clin_intro)   # P019
    _empty()   # P020

    _h3("Patient-level attributes (data_clinical_patient.txt)")   # P021
    _empty()   # P022

    # TABLE 2 — patient column mapping (N x 4c)
    pat_map = _first_colmap(clin_pat) or [
        ("Case ID",       "PATIENT_ID",    "STRING",
         "Required. Use numeric Case ID (e.g., '1', '2')"),
        ("Gender",        "SEX",           "STRING", "Values: Male / Female"),
        ("Primary Location","TUMOR_SITE",  "STRING", "Values: Stomach, Intestine, Other"),
        ("TKI treatment", "TKI_TREATMENT", "STRING", "Values: TKI naive / TKI treated"),
        ("(Not in supp)", "OS_STATUS",     "STRING",
         "\u26a0 MISSING \u2014 needed for survival plots. "
         "Must source from authors or registry. Values: 0:LIVING / 1:DECEASED"),
        ("(Not in supp)", "OS_MONTHS",     "NUMBER",
         "\u26a0 MISSING \u2014 needed for survival plots. "
         "Overall survival in months since diagnosis."),
    ]
    t2 = _newtable(4)
    _hdr_row(t2, ["Source Column (Supp 1)", "cBioPortal Column", "Datatype", "Notes"], W2)
    for i, row in enumerate(pat_map):
        if len(row) == 4:
            src, dst, dtype, note = str(row[0]), str(row[1]), str(row[2]), str(row[3])
        elif len(row) == 3:
            src, dst, note = str(row[0]), str(row[1]), str(row[2]); dtype = "STRING"
        else:
            src, dst, dtype, note = str(row[0]), str(row[1]), "STRING", ""
        _dat_row(t2, [src, dst, dtype, note], W2, row_idx=i)

    _empty()   # P023

    _h3("Sample-level attributes (data_clinical_sample.txt)")   # P024
    _empty()   # P025

    # TABLE 3 — sample column mapping (N x 4c)
    samp_map = _first_colmap(clin_sam) or [
        ("Sample ID",                      "SAMPLE_ID",             "STRING",
         "Required. Use sample IDs as-is (e.g., '1T')"),
        ("Case ID",                        "PATIENT_ID",            "STRING",
         "Required link to patient"),
        ("Sample Type",                    "SAMPLE_TYPE",           "STRING",
         "Fresh Frozen Tissue / Cell Line"),
        ("Clinicopathologic Classification","RISK_STRATIFICATION",  "STRING",
         "Low-risk / Intermediate-risk / High-risk / Metastatic"),
        ("KIT/PDGFRA Genotype",            "DRIVER_MUTATION",       "STRING",
         "Keep as-is; also parse to KIT_EXON and DRIVER_GENE"),
        ("Tumor size (cm)",                "TUMOR_SIZE_CM",         "STRING",
         "Values are ranges (e.g., <=2, 2-5, >5) \u2014 treat as STRING datatype"),
        ("Mit Index",                      "MITOTIC_INDEX",         "STRING",
         "Values: <5 or >=5 per 50 HPF"),
        ("Tumor Rupture (Supp 1)",         "TUMOR_RUPTURE",         "STRING", "Yes / No"),
        ("YLPM1 Mutations (Supp 5)",       "YLPM1_MUTATION_STATUS", "STRING",
         "WT or mutation description"),
        ("YLPM1 Protein (Supp 5)",         "YLPM1_PROTEIN_STATUS",  "STRING",
         "Normal / Loss / NA"),
        ("WES/WGS/WTS flags (Supp 1)",     "SEQUENCING_TYPE",       "STRING",
         "Encode available sequencing: WES, WGS, WTS (comma-separated)"),
    ]
    t3 = _newtable(4)
    _hdr_row(t3, ["Source Column", "cBioPortal Column", "Datatype", "Notes"], W3)
    for i, row in enumerate(samp_map):
        if len(row) == 4:
            src, dst, dtype, note = str(row[0]), str(row[1]), str(row[2]), str(row[3])
        elif len(row) == 3:
            src, dst, note = str(row[0]), str(row[1]), str(row[2]); dtype = "STRING"
        else:
            src, dst, dtype, note = str(row[0]), str(row[1]), "STRING", ""
        _dat_row(t3, [src, dst, dtype, note], W3, row_idx=i)

    _empty()   # P026

    _body("File header format reminder: ",
          "cBioPortal clinical files require 5 header rows before data "
          "(Display Name, Description, Datatype, Priority, Column Name).")   # P027
    _empty()   # P028

    # ── 4.2  MUTATION / MAF ───────────────────────────────────────────────
    maf_src = _files_compact("MUTATION_MAF","MUTATION")
    _h2(f"4.2  Mutation Data \u2014 {maf_src}  \u2192  data_mutations.txt (MAF format)")   # P029

    maf_intro = (_first_intro(maf_recs) or
                 "Somatic coding SNVs and indels. ANNOVAR-annotated and largely maps to "
                 "MAF format, but column remapping is required.")
    _body("", maf_intro)   # P030
    _empty()   # P031

    # TABLE 4 — MAF column mapping (N x 3c)
    maf_map = _first_colmap(maf_recs) or [
        ("Gene",                               "Hugo_Symbol",
         "Verify against HGNC; update outdated gene names"),
        ("Sample ID",                          "Tumor_Sample_Barcode",
         "Must exactly match SAMPLE_ID in clinical file"),
        ("Chromosome",                         "Chromosome",
         "Strip 'chr' prefix (e.g., chr1 \u2192 1)"),
        ("Start Location",                     "Start_Position",
         "Integer \u2014 use as-is"),
        ("End Location",                       "End_Position",
         "Integer \u2014 use as-is"),
        ("Reference",                          "Reference_Allele",
         "Use as-is; '-' for indels is valid"),
        ("Alteration",                         "Tumor_Seq_Allele2",
         "Use as-is; '-' for deletions is valid"),
        ("Mutation Type",                      "Variant_Type",
         "Remap: exonic/splicing context \u2192 SNP / INS / DEL"),
        ("Exonic Mutation Function Annotation","Variant_Classification",
         "Remap ANNOVAR terms (see mapping below)"),
        ("Variant Allele Frequency Mean",      "t_AF (custom)",
         "VAF is available; t_depth/t_alt_count/n_alt_count must be sourced"),
        ("(derive)",                           "NCBI_Build",
         f"Set to GRCh37 (hg19 used for alignment)" if ref_genome == "hg19"
         else "Set to GRCh38 (hg38 used for alignment)"),
        ("(not present)",                      "Matched_Norm_Sample_Barcode",
         "Add matched normal IDs from Supp Data clinical file"),
    ]
    t4 = _newtable(3)
    _hdr_row(t4, ["Supp Data 3 Column", "MAF Column (cBioPortal)", "Transformation Required"], W4)
    for i, row in enumerate(maf_map):
        src = str(row[0]); dst = str(row[1])
        note = str(row[2]) if len(row) > 2 else ""
        _dat_row(t4, [src, dst, note], W4, row_idx=i)

    _empty()   # P032

    _body("ANNOVAR \u2192 MAF Variant_Classification mapping: ",
          "nonsynonymous SNV \u2192 Missense_Mutation  |  "
          "stopgain \u2192 Nonsense_Mutation  |  "
          "stoploss \u2192 Nonstop_Mutation  |  "
          "frameshift deletion \u2192 Frame_Shift_Del  |  "
          "frameshift insertion \u2192 Frame_Shift_Ins  |  "
          "nonframeshift deletion \u2192 In_Frame_Del  |  "
          "nonframeshift insertion \u2192 In_Frame_Ins  |  "
          "splicing \u2192 Splice_Site")   # P033
    _empty()   # P034

    # ── 4.3  STRUCTURAL VARIANTS ──────────────────────────────────────────
    sv_src = _files_compact("STRUCTURAL_VARIANT")
    _h2(f"4.3  Structural Variant Data \u2014 {sv_src}  \u2192  data_sv.txt")   # P035

    sv_intro = (_first_intro(sv_recs) or
                "Structural variants including inversions, deletions, duplications, "
                "and translocations. Core fields present; remap to cBioPortal SV format.")
    _body("", sv_intro)   # P036
    _empty()   # P037

    # TABLE 5 — SV column mapping (N x 3c)
    sv_map = _first_colmap(sv_recs) or [
        ("Sample ID (10a) / Sample ID (14)",     "Sample_Id",
         "Must match SAMPLE_ID in clinical data"),
        ("Left Gene / Gene Involved (10a)",       "Site1_Hugo_Symbol",
         "For 10a: take first gene in 'Gene Involved' (split on '/'). "
         "For 14: use Left Gene directly"),
        ("Right Gene / Gene Involved 2nd (10a)",  "Site2_Hugo_Symbol",
         "For 10a: take second gene. For 14: use Right Gene column"),
        ("Start Chromosome",  "Site1_Chromosome", "Strip 'chr' prefix"),
        ("End Chromosome",    "Site2_Chromosome", "Strip 'chr' prefix"),
        ("Start Position",    "Site1_Position",   "Integer coordinate"),
        ("End Position",      "Site2_Position",   "Integer coordinate"),
        ("SV Type (10a)",     "Class",
         "Remap: INV\u2192INVERSION, DEL\u2192DELETION, "
         "DUP\u2192DUPLICATION, BND\u2192TRANSLOCATION"),
        ("(derive)",          "SV_Status",
         "Set to SOMATIC for all entries (confirmed somatic SVs from WGS/WES)"),
        ("(derive)",          "Site1_Region / Site2_Region",
         "Optional but recommended: EXONIC, INTRONIC, INTERGENIC"),
        ("Arriba reading_frame (14)", "Event_Info",
         "Include fusion reading frame annotation (in-frame vs out-of-frame)"),
    ]
    t5 = _newtable(3)
    _hdr_row(t5, ["Source", "cBioPortal SV Column", "Mapping Instructions"], W5)
    for i, row in enumerate(sv_map):
        src = str(row[0]); dst = str(row[1])
        note = str(row[2]) if len(row) > 2 else ""
        _dat_row(t5, [src, dst, note], W5, row_idx=i)

    _empty()   # P038

    _note("Note: For fusions in Supp Data 14, use the consensus breakpoint position "
          "(present in 2+ callers) as the canonical Site1/Site2 positions. "
          "STAR-Fusion junction coordinates are recommended as they have the highest precision.")   # P039
    _empty()   # P040

    # ── 4.4  COPY NUMBER ──────────────────────────────────────────────────
    # CNA heading: combine arm-level + focal CNA files
    _cna_focal = _files("DISCRETE_CNA","CONTINUOUS_CNA","GISTIC")
    _cna_arm   = [r.get("file","") for r in records if "arm" in r.get("cbio_target_file","").lower() or "arm-level" in (r.get("verdict","") or "").lower()]
    _cna_all   = list(dict.fromkeys(_cna_arm + _cna_focal))
    if _cna_all:
        _pref = next((f[:f.lower().find(" data ")+len(" data ")] for f in _cna_all if " data " in f.lower()), "")
        cna_src = (_pref + " + ".join(f[len(_pref):] for f in _cna_all)) if _pref and all(f.startswith(_pref) for f in _cna_all) else " + ".join(_cna_all)
    else:
        cna_src = "Supp Data"
    _h2(f"4.4  Copy Number Data \u2014 {cna_src}  \u2192  GISTIC / CNA Files")   # P041

    cna_intro = (_first_intro(cna_recs) or
                 "GISTIC 2.0 focal CNV peak output and arm-level frequencies. "
                 "For full cBioPortal GISTIC loading, the complete GISTIC output directory "
                 "is needed (not always available in supplementary). "
                 "However, partial curation is possible.")
    _body("", cna_intro)   # P042
    _empty()   # P043

    # H3 + bullets: focal CNV
    focal_recs = _recs("DISCRETE_CNA","GISTIC")
    focal_src  = focal_recs[0].get("file","Supp Data") if focal_recs else "Supp Data"
    _h3(f"From {focal_src} \u2192 data_gistic_genes_amp.txt / data_gistic_genes_del.txt")  # P044

    focal_bullets = _first_transforms(focal_recs) or [
        "Separate rows by Type (AMP vs DEL) into two files",
        "Required columns: q_value, cytoband, wide_peak_boundaries, gene_list",
        "Gene lists in 'All Genes in Wide Peak' column can be parsed into per-gene rows",
        "Q-values and peak boundaries are already present",
    ]
    for b in focal_bullets[:4]:
        _bullet(str(b))   # P045-P048

    _empty()   # P049

    # H3 + bullets: arm-level CNA
    arm_recs = [r for r in records
                if "arm" in r.get("cbio_target_file","").lower()
                or "arm" in (r.get("verdict","") or "").lower()
                or "arm" in r.get("classification","").lower()]
    arm_src = arm_recs[0].get("file","Supp Data") if arm_recs else "Supp Data"
    _h3(f"From {arm_src} \u2192 Generic Assay: Arm-Level CNA")   # P050

    arm_bullets = _first_transforms(arm_recs) or [
        "cBioPortal supports per-sample arm-level CNA as a Generic Assay",
        "This file has frequency data only \u2014 per-sample arm calls are needed from the original GISTIC run",
        "Contact authors for the GISTIC all_lesions.conf_90 output file which has per-sample calls",
    ]
    for b in arm_bullets[:3]:
        _bullet(str(b))   # P051-P053

    _empty()   # P054

    # ── 4.5  MUTSIG ───────────────────────────────────────────────────────
    ms_src = _files_compact("MUTSIG")
    _h2(f"4.5  MutSig Data \u2014 {ms_src}  \u2192  data_mutsig.txt")   # P055

    ms_intro = (_first_intro(ms_recs) or
                "MutSigCV output for significantly mutated genes. "
                "cBioPortal's MutSig format requires gene, rank, p-value, and q-value.")
    _body("", ms_intro)   # P056
    _empty()   # P057

    # TABLE 6 — MutSig column mapping (N x 3c)
    ms_map = _first_colmap(ms_recs) or [
        ("gene",              "gene",
         "Hugo symbol \u2014 already present"),
        ("codelen (codelenb)","codelen",
         "Coding length \u2014 already present"),
        ("nnon (nnonj)",      "nnon",
         "Non-silent count \u2014 already present"),
        ("(not in sheet)",    "p",
         "\u26a0 P-value column missing \u2014 retrieve from full MutSigCV output (.sig_genes.txt)"),
        ("(not in sheet)",    "q",
         "\u26a0 Q-value/FDR column missing \u2014 retrieve from full MutSigCV output"),
    ]
    t6 = _newtable(3)
    _hdr_row(t6, ["Supp Data 4 Column", "MutSig File Column", "Notes"], W6)
    for i, row in enumerate(ms_map):
        src = str(row[0]); dst = str(row[1])
        note = str(row[2]) if len(row) > 2 else ""
        _dat_row(t6, [src, dst, note], W6, row_idx=i)

    _empty()   # P058

    # ── 4.6  GENE FUSIONS ─────────────────────────────────────────────────
    fus_src = _files_compact("STRUCTURAL_VARIANT") if fus_recs else "Supp Data"
    _h2(f"4.6  Gene Fusions \u2014 {fus_src} (additional notes)")   # P059

    fus_intro = (
        _first_notes(fus_recs) or _first_intro(fus_recs) or
        "Gene fusion events detected by multiple callers. "
        "Only fusions detected by \u22652 callers were retained. Recommended processing:"
    )
    _body("", fus_intro)   # P060

    fus_bullets = _first_transforms(fus_recs) or [
        "Filter for high-confidence fusions: Arriba confidence = 'high' AND STAR-Fusion FFPM > 0.1 threshold",
        "Prioritize in-frame fusions (Arriba reading_frame = 'in-frame') for clinical significance",
        "Notable recurrent fusions should be curated first",
        "For fusions shared with SV file, consolidate into a single SV entry",
    ]
    for b in fus_bullets[:4]:
        _bullet(str(b))   # P061-P064

    _empty()   # P065

    # ──────────────────────────────────────────────────────────────────────
    # §5  CASE LISTS
    # ──────────────────────────────────────────────────────────────────────
    _h1("5. Case Lists")   # P066

    _body("", "cBioPortal requires case lists to define which samples have each data type. "
              "Based on the study, the following case lists should be created:")   # P067
    _empty()   # P068

    # TABLE 7 — case lists (N x 3c)
    has_wes  = any("WES" in str(s).upper() for s in seq_types)
    has_wgs  = any("WGS" in str(s).upper() for s in seq_types)
    has_rna  = any(k in str(s).upper() for s in seq_types
                   for k in ("WTS","RNA","MRNA","RNA-SEQ"))
    has_maf  = bool(maf_recs)
    has_sv   = bool(sv_recs or fus_recs)
    has_cna  = bool(cna_recs)
    has_expr = bool(_recs("EXPRESSION","MRNA"))

    cl_rows: list[tuple] = [
        ("cases_all.txt", f"{study_id}_all",
         f"All {n_samples} samples in the study"),
    ]
    if has_maf or has_wes or has_wgs:
        cl_rows.append(("cases_sequenced.txt", f"{study_id}_sequenced",
                         "Samples with WES or WGS mutation data"))
    if has_cna or has_wes or has_wgs:
        cl_rows.append(("cases_cna.txt", f"{study_id}_cna",
                         "Samples with CNV calls (same as sequenced cohort)"))
    if has_sv:
        cl_rows.append(("cases_sv.txt", f"{study_id}_sv",
                         "Samples with structural variant and/or fusion data"))
    if has_rna or has_expr:
        cl_rows.append(("cases_rna_seq_mrna.txt", f"{study_id}_rna_seq",
                         "Samples with RNA-seq / WTS expression data"))

    t7 = _newtable(3)
    _hdr_row(t7, ["Case List File", "Stable ID", "Samples to Include"], W7)
    for i, (fname, sid, desc) in enumerate(cl_rows):
        _dat_row(t7, [fname, sid, desc], W7, row_idx=i)

    _empty()   # P069

    # ──────────────────────────────────────────────────────────────────────
    # §6  DATA GAPS
    # ──────────────────────────────────────────────────────────────────────
    _h1("6. Data Gaps & Contact with Authors")   # P070

    _body("", "The following data types are referenced in the paper but not available "
              "in the supplementary files. Contacting the corresponding authors is recommended:")   # P071
    _empty()   # P072

    # TABLE 8 — data gaps (N x 2c)
    gap_rows: list[tuple[str,str]] = [
        ("Continuous CNA matrix (log2 ratios)",
         "Per-gene log2 copy number ratios \u2014 needed for data_log2_CNA.txt. "
         "Available from data repository or upon request."),
        ("Discrete CNA matrix (gene-level calls)",
         "Per-gene GISTIC calls (-2/-1/0/1/2) per sample \u2014 needed for data_CNA.txt. "
         "Request from authors or re-run GISTIC."),
        ("RNA-seq expression matrix",
         "Gene-level expression values \u2014 needed for data_mrna_seq_fpkm.txt. "
         "Often available from data repository."),
        ("Segmentation file (CBS segments)",
         "Segmentation output (.seg file) \u2014 needed for genome browser. "
         "Request from authors."),
        ("OS/DFS survival data",
         "Overall and disease-free survival times and status are often not in supplementary. "
         "Essential for Kaplan-Meier plots in cBioPortal Study View."),
        ("Tumor purity / ploidy",
         "Purity/ploidy estimates useful as sample-level clinical attributes "
         "(TUMOR_PURITY, PLOIDY). Request from authors."),
        ("mRNA subtype assignments",
         "Per-sample subtype labels not always in supplementary. "
         "Request from authors or derive from published clustering figure."),
    ]
    # Append study-specific missing items
    seen = {r[0][:40] for r in gap_rows}
    for rec in records:
        for m in (rec.get("missing_required") or []):
            key = str(m)[:40]
            if key not in seen:
                gap_rows.append((str(m)[:60],
                                 "Detected as missing \u2014 source from authors or repository."))
                seen.add(key)
            if len(gap_rows) >= 12:
                break

    t8 = _newtable(2)
    _hdr_row(t8, ["Missing Data", "Details & Recommendation"], W8)
    for i, (lbl, det) in enumerate(gap_rows):
        _dat_row(t8, [lbl, det], W8, row_idx=i)

    _empty()   # P073

    # ──────────────────────────────────────────────────────────────────────
    # §7  CURATION CHECKLIST
    # ──────────────────────────────────────────────────────────────────────
    _h1("7. Curation Checklist & Priority Order")   # P074
    _empty()   # P075

    # TABLE 9 — checklist (N x 4c)
    checklist: list[tuple[str,str,str]] = [
        ("Create meta_study.txt and meta_cancer_type.txt",
         "HIGH", "Use values from Section 2 above"),
    ]
    if clin_pat:
        src = ", ".join(_files("CLINICAL_PATIENT"))
        checklist.append(("Curate data_clinical_patient.txt",
                           "HIGH", f"From {src} \u2014 patient-level rows only"))
    if clin_sam or clin_pat:
        src = ", ".join(dict.fromkeys(_files("CLINICAL_SAMPLE") + _files("CLINICAL_PATIENT")))
        checklist.append(("Curate data_clinical_sample.txt",
                           "HIGH", f"From {src} \u2014 tumor sample rows"))
    if maf_recs:
        src = ", ".join(_files("MUTATION_MAF","MUTATION"))
        checklist.append(("Curate data_mutations.txt (MAF)",
                           "HIGH", f"Remap {src} per Section 4.2 column mapping table"))
    if sv_recs:
        src = ", ".join(_files("STRUCTURAL_VARIANT"))
        checklist.append(("Curate data_sv.txt (DNA SVs)",
                           "HIGH", f"Remap {src} per Section 4.3"))
    if fus_recs:
        src = ", ".join(dict.fromkeys(r.get("file","") for r in fus_recs))
        checklist.append(("Curate data_sv.txt (fusions)",
                           "HIGH", f"Merge {src} into SV file or create data_fusion.txt"))
    checklist.append(("Create case lists",
                       "HIGH", f"{len(cl_rows)} case list files per Section 5"))
    if not cna_recs:
        checklist.append(("Request GISTIC output from authors",
                           "MEDIUM", "Needed for data_CNA.txt, data_log2_CNA.txt, and seg file"))
    if not _recs("EXPRESSION","MRNA"):
        checklist.append(("Request RNA-seq expression matrix",
                           "MEDIUM", "From data repository or corresponding authors"))
    if ms_recs:
        src = ", ".join(_files("MUTSIG"))
        checklist.append(("Curate data_mutsig.txt",
                           "MEDIUM", f"{src} + request p/q-values from authors if missing"))
    checklist.append(("Add arm-level CNA as Generic Assay",
                       "MEDIUM", "Requires per-sample arm calls from GISTIC run"))
    checklist.append(("Add mRNA subtype as clinical attribute",
                       "MEDIUM", "Request per-sample subtype assignments from authors"))
    checklist.append(("Add mutational signatures as Generic Assay",
                       "LOW", "Need per-sample exposure weights from SigProfiler / COSMIC"))
    checklist.append(("Run cBioPortal validateData.py",
                       "HIGH",
                       "Validate all files before loading: "
                       "python validateData.py -s study_dir/ -n"))

    t9 = _newtable(4)
    _hdr_row(t9, ["#", "Task", "Priority", "Source / Action"], W9)
    for i, (task, prio, action) in enumerate(checklist):
        _dat_row(t9, [str(i + 1), task, prio, action], W9, row_idx=i)

    _empty()   # P076

    # ──────────────────────────────────────────────────────────────────────
    # §8  DATA ACCESSION
    # ──────────────────────────────────────────────────────────────────────
    _h1("8. Data Accession & Reproducibility")   # P077

    # Build data repo lines from meta
    repo_labels = ["Raw sequencing data", "Mutation MAF",
                   "External validation cohorts", "Processed data"]
    if data_repos:
        for i, repo in enumerate(data_repos):
            label = repo_labels[i] if i < len(repo_labels) else "Repository"
            _body(f"{label}: ", str(repo))
    else:
        _body("Raw sequencing data: ",
              "Check paper Methods / Data Availability for repository accession numbers.")

    if corresp:
        _body("Corresponding authors: ", corresp)

    _empty()   # P082

    _note("All curation should follow cBioPortal data loading guidelines: "
          "https://docs.cbioportal.org/data-loading/")   # P083

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(output_path)



# ─────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────

def curate(
    pdf_path: str,
    supp_paths: list[str],
    llm_model: str = "openai/gpt-4o",
    temperature: float = 0.2,
    output_path: str | None = None,
) -> dict:
    """
    Run the full cBioPortal curation pipeline.

    Parameters
    ----------
    pdf_path     : path to the main paper PDF
    supp_paths   : list of paths to supplementary Excel files
    llm_model    : model string for load_chat_model (provider/model)
    temperature  : LLM sampling temperature
    output_path  : where to write the .docx report (auto-generated if None)

    Returns
    -------
    {"report_path": str, "summary": {...}}
    """
    # 1. Extract metadata from PDF
    pdf_text = _extract_pdf_text(pdf_path)
    try:
        meta = _extract_metadata_llm(pdf_text, llm_model, temperature)
    except Exception as _llm_exc:
        import logging
        logging.warning(f"LLM metadata extraction failed: {_llm_exc}. Using regex fallback.")
        try:
            meta = _extract_metadata_regex(pdf_text)
        except Exception:
            meta = {
                "study_title":          Path(pdf_path).stem,
                "cancer_type":          "mixed",
                "cancer_type_full":     "Mixed Cancer Type",
                "study_id_suggestion":  "study_upload",
                "description":          "",
                "meta_description":     "",
                "reference_genome":     "hg19",
                "sequencing_types":     [],
                "pmid": "", "doi": "", "year": "", "journal": "",
                "first_author_surname": "", "key_findings": [],
                "primary_site": "", "cohort_description": "",
                "data_repositories": [], "corresponding_authors": "",
            }

    # 2. Analyse supplementary files
    records = _analyse_supplementary_files(supp_paths)

    # 3. Build report
    if output_path is None:
        output_path = os.path.join(
            tempfile.gettempdir(),
            f"cbioportal_curation_{meta.get('study_id_suggestion','report')}.docx"
        )
    _build_report(meta, records, output_path)

    # 4. Return summary
    summary = {
        "study_id":         meta.get("study_id_suggestion"),
        "cancer_type":      meta.get("cancer_type"),
        "num_samples":      meta.get("num_samples"),
        "reference_genome": meta.get("reference_genome"),
        "files_analysed":   len(supp_paths),
        "sheets_analysed":  len(records),
        "high_priority":    sum(1 for r in records if r.get("priority") == "HIGH"),
        "medium_priority":  sum(1 for r in records if r.get("priority") == "MEDIUM"),
        "not_loadable":     sum(1 for r in records if r.get("curability") == "NO"),
        "report_path":      output_path,
        "file_breakdown": [
            {
                "file":           r["file"],
                "sheet":          r["sheet"],
                "cbio_format":    r["cbio_target_file"],
                "curability":     r["curability"],
                "priority":       r["priority"],
                "confidence":     r.get("confidence", 0),
                "verdict":        r.get("verdict", ""),
                "req_present":    r.get("required_present", []),
                "req_missing":    r.get("required_missing", []),
                "opt_present":    r.get("optional_present", []),
            }
            for r in records
        ],
    }
    return {"report_path": output_path, "summary": summary}
